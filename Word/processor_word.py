from docx import Document
from docx.shared import Inches
from visualization_data import read_xlsx, get_text

"""
A function that creates a complete report from the files 'raw_data/sample1.docx' and 'processed_data/processed_data.xlsx'
"""


def data_report():
    doc = Document()

    # Header
    doc.add_heading('Привіт, це аналітика', level=4)

    # Inserting brief information about the sample1.docx file
    docx_file = get_text('../File_System/raw_data/sample1.docx')
    doc.add_paragraph(f'Коротко про цей файл. \n{docx_file[0]}. {docx_file[1]}')

    # Table output
    doc.add_heading('Таблиця всіх присутніх', level=2)

    # Getting data from processed_data.xlsx file
    data_excel = read_xlsx()

    # Creating a table
    table = doc.add_table(rows=len(data_excel) + 1, cols=3)

    # Fill column headers
    table.cell(0, 0).text = 'ID'
    table.cell(0, 1).text = 'Імʼя'
    table.cell(0, 2).text = 'Прізвище'

    # Filling the table with data
    for i, item in enumerate(data_excel):
        id = item.get('id')
        first_name = item.get('first_name')
        last_name = item.get('last_name')
        table.cell(i + 1, 0).text = str(id)
        table.cell(i + 1, 1).text = first_name
        table.cell(i + 1, 2).text = last_name

    # Deleting content (text) from cells
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''
    table.cell(1, 2).text = ''

    # Image output
    doc.add_paragraph('\nЦей графік дасть зрозуміти щось')
    doc.add_picture('media/plot.png', width=Inches(5.0))

    # Save
    doc.save('data_analysis_report.docx')


data_report()
